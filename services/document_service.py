# services/document_service.py
# Document service placeholder.
# Logic for create_docx_entry() and merge_subject_docs() will be implemented here.

import io

from docx.shared import Inches
from pathlib import Path
from datetime import datetime
from docx import Document
from docxcompose.composer import Composer

def create_docx_entry(subject_folder, notes, images, preserve_files) -> str:
    timestamp = datetime.now()
    filename = timestamp.strftime("%d-%m-%Y_%H-%M-%S.docx")

    try:
        subject_folder_path = Path(subject_folder)
        filepath = subject_folder_path/filename

    except Exception as e:
        raise ValueError(
            f"Error creating file path: {e}"
        )

    try:
        document = Document()
    except Exception as e:
        raise ValueError(
            f"Error creating document: {e}"
        )

    document.add_paragraph("=================================")
    document.add_paragraph(timestamp.strftime("%d %b %Y %H:%M"))
    document.add_paragraph("=================================")
    document.add_paragraph("")
    for image in images:
        if image["filename"] not in preserve_files:
            continue

        document.add_heading(
            "Original Images",
            level=2
        )

        image_stream = io.BytesIO(
                image["bytes"]
            )
        try:
            document.add_picture(
                    image_stream,
                    width=Inches(5)
                )
        except Exception as e:
            raise ValueError(
                f"Error adding image to document: {e}"
            )
        document.add_paragraph("")

    document.add_heading(
        "Generated Notes",
        level=2
    )
    document.add_paragraph(notes)
    try:
        document.save(str(filepath))
    except Exception as e:
        raise ValueError(
            f"Error saving document: {e}"
        )

    return str(filepath)


def merge_subject_docs(
    subject_folder: str
) -> str:

    subject_path = Path(subject_folder)

    docx_files = []

    for file in subject_path.glob("*.docx"):

        if file.name.endswith(
            "_Complete_Notes.docx"
        ):
            continue

        if (file.name == "test_merge.docx") or ("_Complete_Notes" in file.name):
            continue

        docx_files.append(file)

    docx_files.sort()

    if len(docx_files) == 0:
        raise ValueError(
            "No notes found for this subject."
        )

    master = Document(
        str(docx_files[0])
    )

    composer = Composer(master)

    for file in docx_files[1:]:

        doc = Document(
            str(file)
        )

        composer.append(doc)

    subject_name = (
        subject_path.name.title()
    )

    output_file = (
        subject_path
        / f"{subject_name}_Complete_Notes.docx"
    )

    composer.save(
        str(output_file)
    )

    return str(output_file)

def get_subject_entries(
    subject_folder: str
):

    entries = []

    for file in Path(subject_folder).glob("*.docx"):

        if file.name.endswith(
            "_Complete_Notes.docx"
        ):
            continue

        entries.append(file.name)

    entries.sort(reverse=True)

    return entries